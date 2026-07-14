"""
EthioBio AI Assistant — Curriculum Ingestion Script (PyMuPDF + OCR)

Scans data/textbooks/ for PDF files organized by grade,
extracts text with PyMuPDF (best quality for Ethiopian curriculum PDFs),
chunks using HybridChunker for token-aware RAG-optimized chunks,
embeds them, and stores in ChromaDB for retrieval.

Usage:
    python scripts/ingest_curriculum.py                               # Ingest all files with PyMuPDF
    python scripts/ingest_curriculum.py --stats                       # Show store stats
    python scripts/ingest_curriculum.py --query "What is cell?"       # Test retrieval
    python scripts/ingest_curriculum.py --clear                       # Clear all vectors
    python scripts/ingest_curriculum.py --use-docling                 # Use Docling instead of PyMuPDF
    python scripts/ingest_curriculum.py --use-easyocr                 # Use EasyOCR (English only, best for garbled PDFs)
"""

import argparse
import glob
import os
import re
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from src.config import settings
from src.rag.embedder import Embedder
from src.rag.vector_store import VectorStore

TEXTBOOKS_DIR = os.path.join(os.path.dirname(__file__), "..", "data", "textbooks")
SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt")


def extract_text_from_pdf(
    filepath: str,
    use_pymupdf: bool = True,
    use_tesseract: bool = False,
    use_easyocr: bool = False,
) -> list[dict]:
    """Extract text from a PDF file.

    Uses PyMuPDF by default (best quality for Ethiopian curriculum PDFs).
    Falls back to Docling+OCR if use_pymupdf=False.
    Uses Tesseract with English+Ethiopic if use_tesseract=True.
    Uses EasyOCR (English only) if use_easyocr=True.
    """
    if use_easyocr:
        return _extract_with_easyocr(filepath)
    if use_tesseract:
        return _extract_with_tesseract(filepath)
    if use_pymupdf:
        return _extract_with_pymupdf(filepath)
    return _extract_with_docling(filepath)


def _extract_with_docling(filepath: str) -> list[dict]:
    """Extract text using PyPdfium2 with RapidOCR fallback for garbled pages.

    If more than 50% of pages are garbled, uses full OCR instead.
    """
    import pypdfium2 as pdfium

    pages = []
    garbled_pages = []

    doc = pdfium.PdfDocument(filepath)
    for i in range(len(doc)):
        page = doc[i]
        text = page.get_textpage().get_text_bounded().strip()
        page_num = i + 1
        if text:
            pages.append({"text": text, "page_number": page_num})
            if _is_garbled(text):
                garbled_pages.append(page_num)
    doc.close()

    # If most pages are garbled, use full OCR instead
    garbled_ratio = len(garbled_pages) / len(pages) if pages else 0
    if garbled_ratio > 0.50:
        print(f"    Most pages garbled ({len(garbled_pages)}/{len(pages)}), using full OCR...")
        return _extract_with_ocr(filepath)

    if garbled_pages:
        print(f"    Garbled pages detected: {garbled_pages}, applying RapidOCR fallback...")
        ocr_pages = _extract_with_ocr(filepath, garbled_pages)
        for ocr_page in ocr_pages:
            for i, p in enumerate(pages):
                if p["page_number"] == ocr_page["page_number"]:
                    pages[i] = ocr_page
                    break

    return pages


def _is_garbled(text: str, alpha_threshold: float = 0.40) -> bool:
    """Detect if text has garbled content based on alphabetic character ratio."""
    if not text or len(text) < 50:
        return False
    alpha = sum(1 for c in text if c.isalpha())
    return (alpha / len(text)) < alpha_threshold


def _contains_control_chars(text: str) -> bool:
    """Detect binary control characters (non-printable, non-whitespace) in text.

    Returns True if the text contains characters likely from PDF encoding issues.
    """
    for c in text:
        cp = ord(c)
        if cp < 0x20 and cp not in (0x09, 0x0A, 0x0C, 0x0D):
            return True
        if 0x7F <= cp <= 0x9F:
            return True
    return False


def _strip_control_chars(text: str) -> str:
    """Remove binary control characters from text."""
    return "".join(c for c in text if c.isprintable() or c in "\n\t\r")


def _extract_with_ocr(filepath: str, page_numbers: list[int] | None = None) -> list[dict]:
    """Extract text using pypdfium2 + RapidOCR (bypasses Docling overhead).

    Renders PDF pages as images and runs OCR on them. Much faster than Docling
    for OCR-only extraction since it avoids loading layout models.
    """
    import pypdfium2 as pdfium
    from rapidocr import RapidOCR

    ocr = RapidOCR()
    doc = pdfium.PdfDocument(filepath)
    pages = []

    for i in range(len(doc)):
        page_num = i + 1
        if page_numbers is not None and page_num not in page_numbers:
            continue

        page = doc[i]
        image = page.render(scale=1.0)  # scale=1.0 is faster, good enough for text
        bitmap = image.to_numpy()

        result = ocr(bitmap)
        if result.txts:
            text = " ".join(result.txts).strip()
            if text:
                pages.append({"text": text, "page_number": page_num})

    doc.close()
    return pages


def _extract_with_pymupdf(filepath: str) -> list[dict]:
    """Extract text using PyMuPDF (fallback)."""
    import fitz
    pages = []
    doc = fitz.open(filepath)
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text().strip()
        if text:
            pages.append({"text": text, "page_number": page_num + 1})
    doc.close()
    return pages


def _extract_with_easyocr(filepath: str, dpi: int = 100) -> list[dict]:
    """Extract text using EasyOCR (English only, no Amharic model available).

    Renders each page at 100 DPI by default (balance of speed vs quality),
    converts to numpy array, and runs EasyOCR readtext with paragraph mode.
    Times on CPU: ~63s/page at 100 DPI, ~135s at 150 DPI.

    EasyOCR 1.7.2 has no Amharic recognition model (character file exists
    but no trained .pth). English-only extraction is high quality.
    """
    import easyocr
    import fitz
    import numpy as np
    from PIL import Image
    reader = easyocr.Reader(["en"], gpu=False, verbose=False)

    doc = fitz.open(filepath)
    total = len(doc)
    results = []

    for i in range(total):
        print(f"    Page {i+1}/{total}...", end=" ", flush=True)
        pix = doc[i].get_pixmap(dpi=dpi)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        img_np = np.array(img)

        text_blocks = reader.readtext(img_np, paragraph=True, detail=0)
        text = " ".join(text_blocks).strip()
        if text:
            results.append({"text": text, "page_number": i + 1})
            print(f"{len(text)} chars", flush=True)
        else:
            print("(empty)", flush=True)

    doc.close()
    return results


def _preprocess_for_tesseract(image):
    """Preprocess a PIL image for better Tesseract accuracy."""
    from PIL import ImageEnhance, ImageFilter

    img = image.convert("L")
    img = ImageEnhance.Contrast(img).enhance(1.5)
    img = img.filter(ImageFilter.SHARPEN)
    img = img.point(lambda x: 0 if x < 128 else 255, "1")
    return img


def _extract_with_tesseract(filepath: str) -> list[dict]:
    """Extract text using Tesseract with English+Ethiopic language support.

    Renders each page at 300 DPI, preprocesses the image, and runs
    Tesseract with LSTM+Legacy engine and English+Ethiopic script.
    Supports mixed English and Amharic text content.
    """
    import pypdfium2 as pdfium
    import pytesseract

    tess_config = "--oem 3 --psm 3 -l eng+script/Ethiopic"
    doc = pdfium.PdfDocument(filepath)
    results = []

    print(f"    Tesseract OCR: {len(doc)} pages at 300 DPI...")
    for i in range(len(doc)):
        page = doc[i]
        bitmap = page.render(scale=300 / 72)
        pil_image = bitmap.to_pil()

        processed = _preprocess_for_tesseract(pil_image)
        text = pytesseract.image_to_string(processed, config=tess_config).strip()

        if text:
            results.append({"text": text, "page_number": i + 1})

    doc.close()
    print(f"    Tesseract complete: {len(results)} pages with text")
    return results
    return pages


def extract_text_from_docx(filepath: str) -> list[dict]:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document
    doc = Document(filepath)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [{"text": full_text, "page_number": 0}]


def detect_source_type(filename: str) -> str:
    """Detect the type of curriculum material from filename."""
    name_lower = filename.lower()
    if "teacher" in name_lower or "guide" in name_lower:
        return "teachers_guide"
    if "workbook" in name_lower or "work" in name_lower:
        return "workbook"
    if "reference" in name_lower or "handbook" in name_lower:
        return "reference"
    if "note" in name_lower or "summary" in name_lower:
        return "teacher_notes"
    return "student_textbook"


def detect_grade_from_path(filepath: str) -> int:
    """Extract grade level from the directory path."""
    match = re.search(r"Grade(\d+)", filepath)
    if match:
        return int(match.group(1))
    return 0


_ROMAN_MAP = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5,
              "VI": 6, "VII": 7, "VIII": 8, "IX": 9, "X": 10}

def _extract_unit(text: str) -> str:
    """Extract unit name from text (e.g. 'Unit 3: Biochemical Molecules', 'Unit I: Sub-fields of Biology')."""
    # Match "Unit N: Title" — capture capitalized title words, allowing hyphens and short prepositions
    match = re.search(r"\bUnit\s+(\d+|[IVXLCDM]+):\s*([A-Z][A-Za-z\-]*(?:\s+(?:[A-Z][A-Za-z\-]*|[a-z]{2,4})){0,8})", text, re.IGNORECASE)
    if match and match.start() < 200:
        num_str = match.group(1)
        name = match.group(2).strip()
        if re.search(r'\d+\.\d+', name) or re.search(r'\b\d{2,}\b', name):
            return ""
        name = re.sub(r'\s+', ' ', name).strip()
        if 4 < len(name) < 80:
            num = _roman_to_int(num_str.upper()) if num_str.isalpha() else int(num_str)
            return f"Unit {num}: {name}" if num else ""
    # Match "Unit Two: Plants" (title-case word number with colon)
    match = re.search(r"\bUnit\s+(One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten)\s*:\s*([A-Z][A-Za-z\-]*(?:\s+(?:[A-Z][A-Za-z\-]*|[a-z]{2,4})){0,8})", text)
    if match and match.start() < 200:
        num_word = match.group(1)
        name = match.group(2).strip()
        if re.search(r'\d+\.\d+', name) or re.search(r'\b\d{2,}\b', name):
            return ""
        name = re.sub(r'\s+', ' ', name).strip()
        if 4 < len(name) < 80:
            num_map = {"ONE": 1, "TWO": 2, "THREE": 3, "FOUR": 4, "FIVE": 5,
                       "SIX": 6, "SEVEN": 7, "EIGHT": 8, "NINE": 9, "TEN": 10}
            num = num_map.get(num_word.upper(), 0)
            return f"Unit {num}: {name}" if num else ""
    # Match "UNIT N WORD" (uppercase word form, e.g. "UNIT FOUR GENETICS")
    match = re.search(r"\bUNIT\s+([A-Z]+)\s+([A-Z][A-Z\s]{3,60})", text)
    if match and match.start() < 200:
        num_word = match.group(1)
        name = match.group(2).strip()
        if name and not re.search(r'\d+\.\d+', name):
            num_map = {"ONE": 1, "TWO": 2, "THREE": 3, "FOUR": 4, "FIVE": 5,
                       "SIX": 6, "SEVEN": 7, "EIGHT": 8, "NINE": 9, "TEN": 10}
            num = num_map.get(num_word.upper(), 0)
            return f"Unit {num}: {name.title()}" if num else ""
    return ""


def _roman_to_int(s: str) -> int:
    """Convert Roman numeral string to integer. Returns 0 on invalid input."""
    values = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100, "D": 500, "M": 1000}
    total = 0
    prev = 0
    for c in reversed(s):
        val = values.get(c, 0)
        if val < prev:
            total -= val
        else:
            total += val
        prev = val
    return total


def _derive_topic_from_unit(unit_str: str) -> str:
    """Extract topic from unit string, e.g. 'Unit 3: Biochemical Molecules' -> 'Biochemical Molecules'."""
    if ":" in unit_str:
        return unit_str.split(":", 1)[1].strip()
    return unit_str


def _extract_topic(text: str) -> str:
    """Extract topic/section heading from text."""
    patterns = [
        r"(?:^|\n)\s*(\d+\.\d+\s+[A-Z][^\n]{3,80})",
        r"(?:^|\n)\s*(Topic\s+\d+[:.\s]+[A-Z][^\n]{3,80})",
        r"(?:^|\n)\s*(Lesson\s+\d+[:.\s]+[A-Z][^\n]{3,80})",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()
    return ""


_SECTION_RE = re.compile(r"(?:^|\n)\s*(\d+\.\d+)(?!\.\d)\s+([A-Z][^\n]{3,80})")
_SUBTOPIC_RE = re.compile(r"(?:^|\n)\s*(\d+\.\d+\.\d+)\s+([A-Z][^\n]{3,80})")


def _extract_heading_info(text: str) -> tuple[str, str]:
    """Extract (section, subtopic) from text.

    Section examples: '3.1 Carbohydrates'
    Subtopic examples: '3.1.1 Monosaccharides'
    Returns (heading, '') for sections, ('', heading) for subtopics.
    """
    m = _SUBTOPIC_RE.search(text)
    if m:
        return "", f"{m.group(1)} {m.group(2).strip()}"
    m = _SECTION_RE.search(text)
    if m:
        return f"{m.group(1)} {m.group(2).strip()}", ""
    return "", ""


def _extract_page_number(page_text: str, pdf_page_num: int, grade: int) -> int:
    """Extract the printed textbook page number from page footer or estimate from PDF index.

    Tries common Ethiopian textbook footer patterns first;
    falls back to PDF index - 3 (covers cover/TOC pages before content).
    """
    lines = page_text.strip().split('\n')

    # Pattern 1: standalone number in the last few lines — most reliable
    for line in reversed(lines[-3:]):
        line = line.strip()
        if re.match(r'^\d{1,3}$', line):
            n = int(line)
            if 1 <= n <= 600:
                return n

    footer_region = '\n'.join(lines[-3:]).strip()

    # Pattern 2: "Grade X Biology N" (number after grade/subject)
    m = re.search(rf'Grade\s*{grade}\s*Biology[^A-Za-z0-9]*(\d+)', footer_region)
    if m:
        return int(m.group(1))

    # Pattern 3: "N Grade X Biology" or "N | Grade X Biology" (number before grade/subject)
    m = re.search(rf'(\d+)\s*[|\u2013\u2014\-]?\s*Grade\s*{grade}\s*Biology', footer_region)
    if m:
        return int(m.group(1))

    return max(1, pdf_page_num - 3)


def chunk_text(text: str, source_type: str = "student_textbook") -> list[dict]:
    """Split text into curriculum-aligned chunks at natural boundaries.

    Strategy:
    1. Split on chapter/unit/section headers
    2. Split long sections by paragraphs
    3. Keep chunks between 300-1500 characters
    4. Track unit, topic, and page_number for each chunk
    """
    # No TOC skipping needed — the unit extraction rejects TOC entries
    # (they contain section refs like "1.1" which fail the clean name check)

    header_patterns = [
        r"(?=(?:^|\n)\s*(?:Chapter|Unit|Module)\s+\w+[\.\s:])",
        r"(?=(?:^|\n)\s*\d+(?:\.\d+)+\s+[A-Z])",
        r"(?=(?:^|\n)\s*[A-Z][A-Z\s]{5,}(?:\n|$))",
        r"(?=(?:^|\n)\s*Lesson\s+\d+)",
        r"(?=(?:^|\n)\s*Topic\s+\d+)",
        r"(?=(?:^|\n)\s*Session\s+\d+)",
        r"(?=(?:^|\n)\s*Part\s+[A-Z0-9])",
        r"(?=(?:^|\n)\s*Section\s+\d+)",
    ]

    combined_pattern = "|".join(header_patterns)
    chunks = []

    if source_type == "teachers_guide" or source_type == "teacher_notes":
        paragraphs = re.split(r"\n\s*\n", text)
        current_chunk = ""
        current_unit = ""
        current_section = ""
        current_subtopic = ""
        current_page = 0
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            unit_match = re.search(r"Unit\s+(\d+|[IVXLCDM]+):\s*([A-Z][^\n]{3,60})", para)
            if unit_match:
                name = unit_match.group(2).strip()
                if not re.search(r'\d+\.\d+', name) and not re.search(r'\b\d{2,}\b', name):
                    current_unit = f"Unit {unit_match.group(1)}: {name}"
                    current_section = ""
                    current_subtopic = ""

            sec, sub = _extract_heading_info(para)
            if sec:
                current_section = sec
                current_subtopic = ""
            if sub:
                current_subtopic = sub

            page_match = re.search(r"\[PAGE\s+(\d+)\]", para)
            if page_match:
                current_page = int(page_match.group(1))
            else:
                page_match = re.search(r"Grade\s+\d+\s+Biology\s+(\d+)", para)
                if page_match:
                    current_page = int(page_match.group(1))

            if len(current_chunk) + len(para) > 1500 and current_chunk:
                chunks.append({
                    "text": current_chunk.strip(),
                    "heading": _extract_heading(current_chunk),
                    "unit": current_unit,
                    "section": current_section,
                    "subtopic": current_subtopic,
                    "topic": _extract_topic(current_chunk),
                    "page_number": current_page,
                })
                current_chunk = para
            else:
                current_chunk += "\n\n" + para if current_chunk else para

        if current_chunk:
            chunks.append({
                "text": current_chunk.strip(),
                "heading": _extract_heading(current_chunk),
                "unit": current_unit,
                "section": current_section,
                "subtopic": current_subtopic,
                "topic": _extract_topic(current_chunk),
                "page_number": current_page,
            })
    else:
        sections = re.split(combined_pattern, text)
        current_unit = ""
        current_section = ""
        current_subtopic = ""
        current_page = 0

        for section in sections:
            section = section.strip()
            if not section or len(section) < 30:
                continue

            # Extract unit from this section's header
            unit_match = re.search(r"(?:^|\n)\s*Unit\s+(\d+|[IVXLCDM]+|One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten):\s*([A-Z][^\n\t]{3,60})", section)
            if unit_match:
                name = unit_match.group(2).strip()
                # Reject TOC entries: contain section refs, page numbers, or
                # are followed by mostly whitespace/numbers
                if not re.search(r'\d+\.\d+', name) and not re.search(r'\b\d{2,}\b', name):
                    # Check if the section after the unit header has real content (not just TOC)
                    after_unit = section[unit_match.end():unit_match.end()+100].strip()
                    # TOC entries are followed by whitespace and page numbers, not paragraphs
                    if after_unit and not re.match(r'^[\s\d\t]+$', after_unit[:50]):
                        name = re.split(r'\s{2,}|\t|•', name)[0].strip()
                        if len(name) < 80 and len(name) > 3:
                            raw = unit_match.group(1)
                            _word_to_num = {"One": 1, "Two": 2, "Three": 3, "Four": 4, "Five": 5,
                                            "Six": 6, "Seven": 7, "Eight": 8, "Nine": 9, "Ten": 10}
                            if raw.isdigit():
                                unit_num = int(raw)
                            elif raw in _word_to_num:
                                unit_num = _word_to_num[raw]
                            else:
                                unit_num = _roman_to_int(raw.upper()) if raw.isalpha() else 0
                            current_unit = f"Unit {unit_num}: {name}"
                            current_section = ""
                            current_subtopic = ""

            # Extract section/subtopic from this section's header
            sec, sub = _extract_heading_info(section)
            if sec:
                current_section = sec
                current_subtopic = ""
            if sub:
                current_subtopic = sub

            # Extract page number from marker or text
            page_match = re.search(r"\[PAGE\s+(\d+)\]", section)
            if page_match:
                current_page = int(page_match.group(1))
            else:
                page_match = re.search(r"Grade\s+\d+\s+Biology\s+(\d+)", section)
                if page_match:
                    current_page = int(page_match.group(1))

            if len(section) > 2000:
                paragraphs = re.split(r"\n\s*\n", section)
                current = ""
                for para in paragraphs:
                    para = para.strip()
                    if not para:
                        continue

                    p_match = re.search(r"\[PAGE\s+(\d+)\]", para)
                    if p_match:
                        current_page = int(p_match.group(1))
                    else:
                        p_match = re.search(r"Grade\s+\d+\s+Biology\s+(\d+)", para)
                        if p_match:
                            current_page = int(p_match.group(1))

                    if len(current) + len(para) > 1200 and current:
                        chunks.append({
                            "text": current.strip(),
                            "heading": _extract_heading(current),
                            "unit": current_unit,
                            "section": current_section,
                            "subtopic": current_subtopic,
                            "topic": _extract_topic(current),
                            "page_number": current_page,
                        })
                        current = para
                    else:
                        current += "\n\n" + para if current else para
                if current:
                    chunks.append({
                        "text": current.strip(),
                        "heading": _extract_heading(current),
                        "unit": current_unit,
                        "section": current_section,
                        "subtopic": current_subtopic,
                        "topic": _extract_topic(current),
                        "page_number": current_page,
                    })
            else:
                section_unit = current_unit
                if not section_unit:
                    section_unit = _extract_unit(section)
                chunks.append({
                    "text": section,
                    "heading": _extract_heading(section),
                    "unit": section_unit,
                    "section": current_section,
                    "subtopic": current_subtopic,
                    "topic": _extract_topic(section),
                    "page_number": current_page,
                })

    filtered = [c for c in chunks if len(c["text"]) >= 50]
    return filtered


def _extract_heading(text: str) -> str:
    """Extract the first line that looks like a heading."""
    lines = text.strip().split("\n")
    for line in lines[:5]:
        line = line.strip()
        if line and (line.isupper() or re.match(r"^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*:", line)):
            return line[:100]
    return lines[0][:100] if lines else ""


def scan_files(base_dir: str) -> list[dict]:
    """Scan the textbooks directory for curriculum files organized by grade."""
    files = []
    for grade_dir in sorted(glob.glob(os.path.join(base_dir, "Grade*"))):
        grade = detect_grade_from_path(grade_dir)
        if not grade:
            continue
        for ext in SUPPORTED_EXTENSIONS:
            for fp in sorted(glob.glob(os.path.join(grade_dir, f"*{ext}"))):
                filename = os.path.basename(fp)
                files.append({
                    "filepath": fp,
                    "filename": filename,
                    "grade_level": grade,
                    "source_type": detect_source_type(filename),
                    "extension": ext,
                })
    return files


async def process_file(
    file_info: dict,
    embedder: Embedder,
    store: VectorStore,
    use_pymupdf: bool = True,
    use_tesseract: bool = False,
    use_easyocr: bool = False,
) -> int:
    """Process a single curriculum file: extract, chunk, embed, store."""
    filepath = file_info["filepath"]
    grade = file_info["grade_level"]
    source_type = file_info["source_type"]
    filename = file_info["filename"]

    if use_easyocr:
        extractor = "EasyOCR"
    elif use_tesseract:
        extractor = "Tesseract+Ethiopic"
    elif use_pymupdf and filepath.endswith(".pdf"):
        extractor = "PyMuPDF"
    else:
        extractor = "Docling+OCR"
    print(f"  Processing: Grade {grade}/{filename} ({source_type}) [{extractor}]")

    try:
        if filepath.endswith(".pdf"):
            pages = extract_text_from_pdf(filepath, use_pymupdf=use_pymupdf, use_tesseract=use_tesseract, use_easyocr=use_easyocr)
        elif filepath.endswith(".docx"):
            pages = extract_text_from_docx(filepath)
        elif filepath.endswith(".txt"):
            with open(filepath, "r", encoding="utf-8") as f:
                pages = [{"text": f.read(), "page_number": 0}]
        else:
            return 0
    except Exception as e:
        print(f"    Extraction error: {e}")
        return 0

    if not pages or not any(p["text"] for p in pages):
        print("    ⚠️  No text extracted")
        return 0

    # Chunk per-page to preserve page numbers, then sub-chunk if needed
    all_chunks = []
    for p in pages:
        cleaned_text = _strip_control_chars(p["text"])
        page_chunks = chunk_text(cleaned_text, source_type)
        for c in page_chunks:
            c["page_number"] = _extract_page_number(p["text"], p["page_number"], grade)
            all_chunks.append(c)
    chunks = all_chunks

    if not chunks:
        print("    ⚠️  No chunks created")
        return 0

    # Filter out garbled chunks: too short, control chars, or low alpha ratio
    before = len(chunks)
    filtered_chunks = []
    for c in chunks:
        text = c["text"]
        if len(text) < 80:
            continue
        if _is_garbled(text, alpha_threshold=0.50):
            continue
        if _contains_control_chars(text):
            continue
        filtered_chunks.append(c)
    chunks = filtered_chunks
    if before != len(chunks):
        print(f"    Filtered out {before - len(chunks)} garbled chunks")

    if not chunks:
        print("    ⚠️  All chunks garbled, skipping file")
        return 0

    print(f"    Extracted {len(pages)} pages -> {len(chunks)} quality chunks")

    # Post-chunk metadata extraction fallback: extract unit from chunk body
    _SECTION_UNIT_MAP = {
        "1": "Unit 1: Sub-Fields of Biology",
        "2": "Unit 2: Plants",
        "3": "Unit 3: Biochemical Molecules",
        "4": "Unit 4: Cell Cycle",
        "5": "Unit 5: Human Biology",
        "6": "Unit 6: Ecological Interactions",
    }
    for chunk in chunks:
        if not chunk.get("unit"):
            chunk["unit"] = _extract_unit(chunk["text"])
        if not chunk.get("unit") and chunk.get("section"):
            sec_match = re.match(r"^(\d+)", chunk["section"])
            if sec_match and sec_match.group(1) in _SECTION_UNIT_MAP:
                chunk["unit"] = _SECTION_UNIT_MAP[sec_match.group(1)]
        if not chunk.get("unit"):
            sec_match = re.search(r"(?:^|\n)\s*(\d+)\.\d", chunk["text"])
            if sec_match and sec_match.group(1) in _SECTION_UNIT_MAP:
                chunk["unit"] = _SECTION_UNIT_MAP[sec_match.group(1)]
        if not chunk.get("heading"):
            chunk["heading"] = _extract_heading(chunk["text"])

    chunk_texts = [c["text"] for c in chunks]
    metadatas = []
    ids = []

    for i, chunk in enumerate(chunks):
        chunk_id = f"g{grade}_{filename}_{i}"
        metadatas.append({
            "grade_level": grade,
            "source_type": source_type,
            "source_file": filename,
            "unit": chunk.get("unit", "") or "",
            "section": chunk.get("section", "") or "",
            "subtopic": chunk.get("subtopic", "") or "",
            "topic": chunk.get("topic", "") or _derive_topic_from_unit(chunk.get("unit", "")),
            "heading": chunk.get("heading", "") or chunk["text"][:80],
            "page_number": chunk.get("page_number", 0) or 0,
            "chunk_index": i,
        })
        ids.append(chunk_id)

    try:
        embeddings = await embedder.embed_batch(chunk_texts)
        await store.add_documents(
            texts=chunk_texts,
            embeddings=embeddings,
            metadatas=metadatas,
            ids=ids,
        )
        print(f"    Stored {len(chunks)} chunks in ChromaDB")
        return len(chunks)
    except Exception as e:
        print(f"    Embedding/storage error: {e}")
        return 0


async def main():
    parser = argparse.ArgumentParser(
        description="Ingest curriculum materials into RAG vector store"
    )
    parser.add_argument(
        "--textbooks-dir", default=TEXTBOOKS_DIR,
        help="Path to textbooks directory"
    )
    parser.add_argument("--clear", action="store_true", help="Clear all vectors before ingestion")
    parser.add_argument("--stats", action="store_true", help="Show store statistics")
    parser.add_argument("--query", type=str, help="Test retrieval with a query string")
    parser.add_argument("--grade", type=int, help="Filter query by grade level")
    parser.add_argument(
        "--use-docling", action="store_true",
        help="Use Docling+OCR instead of PyMuPDF"
    )
    parser.add_argument(
        "--use-tesseract", action="store_true",
        help="Use Tesseract OCR with English+Ethiopic support (best for mixed Amharic/English PDFs)"
    )
    parser.add_argument(
        "--use-easyocr", action="store_true",
        help="Use EasyOCR (English only, no Amharic model). Best for garbled English PDFs."
    )
    parser.add_argument(
        "--ollama-embed", action="store_true",
        help="Use Ollama for embeddings instead of local model"
    )

    args = parser.parse_args()

    if not os.path.isdir(args.textbooks_dir):
        print(f"Textbooks directory not found: {args.textbooks_dir}")
        print(f"   Create it and place PDFs in: {args.textbooks_dir}/Grade7/ ... Grade12/")
        sys.exit(1)

    embedder = Embedder()
    store = VectorStore()

    if args.clear:
        if args.grade:
            print(f"Clearing vectors for Grade {args.grade} only...")
            coll = store._get_collection()
            coll.delete(where={"grade_level": args.grade})
            print("   Cleared.")
        else:
            print("Clearing all existing vectors...")
            await store.delete_collection()
            print("   Cleared.")

    if args.stats:
        count = store.count()
        print("📊 Vector Store Statistics")
        print(f"   Collection: {settings.collection_name}")
        print(f"   Total chunks: {count}")
        print(f"   Persist path: {settings.vector_store_path}")
        return

    if args.query:
        from src.rag.retriever import Retriever
        print(f"Testing retrieval for: \"{args.query}\"")
        retriever = Retriever(embedder=embedder, vector_store=store)
        results = await retriever.retrieve(
            query=args.query,
            n_results=3,
            grade_level=args.grade,
        )
        print(f"   Retrieved {len(results)} results")
        for i, r in enumerate(results, 1):
            meta = r.get("metadata", {})
            src = meta.get('source_file', '?')
            print(f"\n   [{i}] Grade {meta.get('grade_level', '?')} - {src}")
            print(f"       {r['content'][:150]}...")
            print(f"       Score: {r.get('score', 0.0):.3f}")
        return

    files = scan_files(args.textbooks_dir)

    if args.grade:
        files = [f for f in files if f["grade_level"] == args.grade]
        print(f"Filtered to Grade {args.grade} only")

    if not files:
        print(f"No supported files found in {args.textbooks_dir}")
        print(f"   Supported: {', '.join(SUPPORTED_EXTENSIONS)}")
        print("   Expected structure:")
        print(f"     {args.textbooks_dir}/Grade7/biology_textbook.pdf")
        print(f"     {args.textbooks_dir}/Grade9/biology_teachers_guide.docx")
        sys.exit(0)

    use_pymupdf = not (args.use_docling or args.use_tesseract or args.use_easyocr)
    use_tesseract = args.use_tesseract
    use_easyocr = args.use_easyocr
    if use_easyocr:
        label = "EasyOCR"
    elif use_tesseract:
        label = "Tesseract+Ethiopic"
    elif use_pymupdf:
        label = "PyMuPDF"
    else:
        label = "Docling+OCR"
    print(f"Found {len(files)} files to process (extractor: {label})\n")
    total_chunks = 0
    for f in files:
        chunks = await process_file(f, embedder, store, use_pymupdf=use_pymupdf, use_tesseract=use_tesseract, use_easyocr=use_easyocr)
        total_chunks += chunks

    # Rebuild BM25 index for hybrid search
    print("\nRebuilding BM25 index...")
    from src.retrieval.adapter import VectorStoreAdapter
    adapter = VectorStoreAdapter(vector_store=store)
    adapter.build_bm25_index()
    print("   BM25 index rebuilt")

    count = store.count()
    print("\n✅ Ingestion complete!")
    print(f"   Files processed: {len(files)}")
    print(f"   Total chunks stored: {total_chunks}")
    print(f"   ChromaDB collection count: {count}")
    print(f"   Location: {settings.vector_store_path}")


if __name__ == "__main__":
    import asyncio
    asyncio.run(main())
