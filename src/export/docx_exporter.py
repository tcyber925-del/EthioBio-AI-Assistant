import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def export_quiz_to_docx(quiz: dict[str, Any], questions: list[dict[str, Any]]) -> bytes:
    doc = Document()

    title = doc.add_heading(quiz.get("title", "Quiz"), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Grade {quiz.get('grade_level', 'N/A')} - {quiz.get('topic', 'N/A')}")
    run.font.size = Pt(12)
    run.font.color.rgb = None

    doc.add_paragraph()

    for i, q in enumerate(questions, 1):
        q_type = q.get("question_type", "unknown").replace("_", " ").title()
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ({q_type}) {q['question_text']}")
        run.bold = True
        run.font.size = Pt(11)

        options = q.get("options")
        if options:
            for opt in options:
                doc.add_paragraph(opt, style="List Bullet")

        p_ans = doc.add_paragraph()
        run_ans = p_ans.add_run(f"Answer: {q['correct_answer']}")
        run_ans.font.size = Pt(10)
        run_ans.font.color.rgb = None

        explanation = q.get("explanation")
        if explanation:
            p_exp = doc.add_paragraph()
            run_exp = p_exp.add_run(f"Explanation: {explanation}")
            run_exp.font.size = Pt(10)
            run_exp.italic = True

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def export_lesson_plan_to_docx(lesson: dict[str, Any]) -> bytes:
    doc = Document()

    title = doc.add_heading("Lesson Plan", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Grade {lesson.get('grade_level', 'N/A')} - {lesson.get('topic', 'N/A')}")
    run.font.size = Pt(12)

    doc.add_paragraph()

    sections = [
        ("Objective", lesson.get("objective", "")),
        ("Prior Knowledge", lesson.get("prior_knowledge", "")),
        ("Explanation", lesson.get("explanation", "")),
        ("Activities", lesson.get("activities", [])),
        ("Assessment", lesson.get("assessment", "")),
        ("Homework", lesson.get("homework", "")),
        ("Teacher Notes", lesson.get("teacher_notes", "")),
    ]

    for section_name, content in sections:
        doc.add_heading(section_name, level=2)
        if isinstance(content, list):
            for item in content:
                if isinstance(item, dict):
                    for key, val in item.items():
                        p = doc.add_paragraph()
                        run = p.add_run(f"{key.replace('_', ' ').title()}: ")
                        run.bold = True
                        p.add_run(str(val))
                else:
                    doc.add_paragraph(str(item), style="List Bullet")
        elif content:
            doc.add_paragraph(content)

    periods = lesson.get("periods")
    if periods:
        doc.add_heading("Lesson Periods", level=2)
        for period in periods:
            p_name = period.get("name", "")
            p_dur = period.get("duration_minutes", "")
            head_p = doc.add_paragraph()
            run = head_p.add_run(f"{p_name} ({p_dur} min)")
            run.bold = True
            run.font.size = Pt(11)
            p_obj = period.get("objective")
            if p_obj:
                obj_p = doc.add_paragraph()
                r_obj = obj_p.add_run("Objective: ")
                r_obj.bold = True
                obj_p.add_run(p_obj)
            p_desc = period.get("description", "")
            if p_desc:
                doc.add_paragraph(p_desc)
            p_ta = period.get("teacher_activity")
            if p_ta:
                ta_p = doc.add_paragraph()
                r_ta = ta_p.add_run("Teacher: ")
                r_ta.bold = True
                ta_p.add_run(p_ta)
            p_sa = period.get("student_activity")
            if p_sa:
                sa_p = doc.add_paragraph()
                r_sa = sa_p.add_run("Students: ")
                r_sa.bold = True
                sa_p.add_run(p_sa)
            p_mat = period.get("materials_needed")
            if p_mat:
                mat_p = doc.add_paragraph()
                r_mat = mat_p.add_run("Materials: ")
                r_mat.bold = True
                mat_p.add_run(", ".join(p_mat) if isinstance(p_mat, list) else str(p_mat))

    exit_ticket = lesson.get("exit_ticket", [])
    if exit_ticket:
        doc.add_heading("Exit Ticket", level=2)
        for i, q in enumerate(exit_ticket, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"Q{i}. ({q.get('question_type', '')}) {q.get('question_text', '')}")
            run.bold = True
            run.font.size = Pt(11)
            options = q.get("options")
            if options:
                for opt in options:
                    doc.add_paragraph(opt, style="List Bullet")
            p_ans = doc.add_paragraph()
            run_ans = p_ans.add_run(f"Answer: {q.get('correct_answer', '')}")
            run_ans.font.size = Pt(10)
            exp = q.get("explanation")
            if exp:
                p_exp = doc.add_paragraph()
                run_exp = p_exp.add_run(f"Explanation: {exp}")
                run_exp.font.size = Pt(10)
                run_exp.italic = True

    differentiation = lesson.get("differentiation", [])
    if differentiation:
        doc.add_heading("Differentiated Activities", level=2)
        for d in differentiation:
            p = doc.add_paragraph()
            run = p.add_run(f"[{d.get('group', '').title()}] ")
            run.bold = True
            p.add_run(f"{d.get('description', '')} ({d.get('duration_minutes', '')}min)")

    diagram_suggestions = lesson.get("diagram_suggestions", [])
    if diagram_suggestions:
        doc.add_heading("Diagram Suggestions", level=2)
        for d in diagram_suggestions:
            p = doc.add_paragraph()
            run = p.add_run(f"{d.get('title', '')} ")
            run.bold = True
            p.add_run(f"({d.get('diagram_type', '')})")
            desc = d.get("description")
            if desc:
                doc.add_paragraph(desc)

    misconception_activities = lesson.get("misconception_activities", [])
    if misconception_activities:
        doc.add_heading("Misconception Activities", level=2)
        for a in misconception_activities:
            p = doc.add_paragraph()
            run = p.add_run(f"{a.get('activity_name', '')} ")
            run.bold = True
            p.add_run(f"({a.get('activity_type', '')} · {a.get('duration_minutes', '')}min)")
            p2 = doc.add_paragraph()
            p2.add_run(f"Addressing: {a.get('misconception', '')}")
            desc = a.get("description")
            if desc:
                doc.add_paragraph(desc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
