import io
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fpdf import FPDF

from src.utils.svg_render import render_svg_to_png


def _svg_to_png_bytes(svg: str, width: int = 800, height: int = 600) -> Optional[bytes]:
    try:
        return render_svg_to_png(svg, width=width, height=height)
    except Exception:
        return None


def export_diagram_to_docx(
    svg: str,
    title: str = "Science Diagram",
    topic: str = "",
    grade: int = 10,
    labels: Optional[list[dict]] = None,
) -> bytes:
    doc = Document()

    title_heading = doc.add_heading(title, level=1)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Grade {grade} - {topic}" if topic else f"Grade {grade}")
    run.font.size = Pt(12)

    doc.add_paragraph()

    png_data = _svg_to_png_bytes(svg)
    if png_data:
        buf = io.BytesIO(png_data)
        doc.add_picture(buf, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[SVG rendering failed — diagram not available]")

    doc.add_paragraph()

    if labels:
        doc.add_heading("Labels", level=2)
        for i, label in enumerate(labels, 1):
            text = label.get("text", "")
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {text}")
            run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


class DiagramPDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 10)
        self.cell(0, 8, "EthioSci", align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def export_diagram_to_pdf(
    svg: str,
    title: str = "Science Diagram",
    topic: str = "",
    grade: int = 10,
    labels: Optional[list[dict]] = None,
) -> bytes:
    pdf = DiagramPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 14, title, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    if topic:
        pdf.set_font("Helvetica", "", 11)
        pdf.cell(0, 8, f"Grade {grade} - {topic}", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    png_data = _svg_to_png_bytes(svg)
    if png_data:
        img_buf = io.BytesIO(png_data)
        pdf.image(img_buf, x=10, w=pdf.w - 20)
    else:
        pdf.set_font("Helvetica", "I", 10)
        pdf.cell(0, 10, "SVG rendering failed", new_x="LMARGIN", new_y="NEXT")

    if labels:
        pdf.ln(6)
        if pdf.get_y() > 240:
            pdf.add_page()
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 8, "Labels", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 10)
        for i, label in enumerate(labels, 1):
            if pdf.get_y() > 270:
                pdf.add_page()
            text = label.get("text", "")
            pdf.cell(0, 5, f"{i}. {text}", new_x="LMARGIN", new_y="NEXT")

    return bytes(pdf.output())
